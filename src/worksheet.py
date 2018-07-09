from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from src.pos import POS
from src.api.nlp import parse_all
import json

PUNCT = {',', '.', '-', "'s", "'m", '?', "n't"}


class Worksheet:
    def __init__(self, lines=[], title='Sentence Worksheet', loc='Worksheet.docx', sources=[], key=False, font_size=13):
        self.key = key
        self.font_size = font_size
        if key:
            index = loc.rindex('.')
            self.loc = loc[:index] + '(Key)' + loc[index:]
            self.title = title + ' (Key)'
        else:
            self.loc = loc
            self.title = title
        self.font = 'Times New Roman'
        self.styles = load_json('config/pos.json')
        self.lines = parse_all(lines)
        self.sources = sources
        self.doc = Document()

    def render(self):
        self.doc = Document()
        self.add_title(self.title)
        if len(self.sources) > 0:
            self.add_sources()
        self.add_instructions()
        self.add_title('')
        for line in self.lines:
            self.add_line(line)
        self.format_document()
        self.doc.save(self.loc)

    def add_title(self, text):
        title = self.doc.add_paragraph(text)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        font = title.style.font
        font.name = self.font
        font.size = Pt(16)
        font.bold = True

    def add_sources(self):
        p = self.doc.add_paragraph()
        self.format_run(p.add_run('Sources: '), font_size=8, bold=True)
        self.format_run(p.add_run(', '.join(self.sources)), font_size=8, bold=False)

    def add_instructions(self):
        self.add_instruction_with_colors()
        self.add_instruction('Insert ', 'any needed commas, and circle them')
        
    def add_instruction_with_colors(self):
        line = self.doc.add_paragraph()
        subtitle = line.add_run('Label: ')
        self.format_run(subtitle, font_size=self.font_size, bold=True)
        comma = False
        for style in self.styles.values():
            if style['active']:
                if comma:
                    self.format_run(line.add_run(", "), font_size=self.font_size)
                else:
                    comma = True
                self.format_run(line.add_run(style['name']), font_size=self.font_size, color=load_color(style['rgb']))

    def add_instruction(self, subtitle, instruction):
        line = self.doc.add_paragraph()
        subtitle = line.add_run(subtitle)
        self.format_run(subtitle, font_size=self.font_size, bold=True)
        self.format_run(line.add_run(instruction), font_size=self.font_size)
        
    def add_line(self, line):
        paragraph = self.doc.add_paragraph(style='List Number')
        self.format_paragraph(paragraph)
        run = None
        current_prep = -1
        if self.key:
            pos = line['pos']
        else:
            pos = [None] * len(line['pos'])
        for index, (word, color) in enumerate(zip(line['doc'], pos)):
            if run is not None and word not in PUNCT:
                run = paragraph.add_run(' ')
                self.format_run(run)
            if type(color) == int and self.styles['Preposition']['active']:
                if index == current_prep:
                    run = paragraph.add_run(word + ')')
                elif index > current_prep:
                    run = paragraph.add_run('(' + word)
                    current_prep = rindex(pos, color)
                else:
                    run = paragraph.add_run(word)
            else:
                run = paragraph.add_run(word)
            self.format_run(run, color=color)

        if str(line['doc'][-1]) not in PUNCT:
            paragraph.add_run('.')

    def format_document(self):
        for section in self.doc.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

    def format_paragraph(self, p):
        p.style.font.bold = False
        p.style.font.size = Pt(self.font_size)
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style = p.style.paragraph_format
        style.line_spacing = Pt(self.font_size * 3)

    def format_run(self, run, color=None, font_size=13, bold=False):
        if type(color) == POS:
            key = list(self.styles.keys())[color.value]
            if self.styles[key]['active']:
                run.font.color.rgb = load_color(self.styles[key]['rgb'])
        elif type(color) == RGBColor:
            run.font.color.rgb = color
        run.bold = bold
        run.font.name = self.font
        run.font.size = Pt(font_size)
        return run


def rindex(lst, val):
    return len(lst) - 1 - lst[::-1].index(val)


def load_color(rgb):
    return RGBColor(*rgb)


def load_json(loc):
    with open(loc) as f:
        return json.load(f)


if __name__ == '__main__':
    sheet = Worksheet(['In the summer, extreme temperatures of over 100 degrees can give visitors heat strokes.',
                       'Big Bend National Park is one of only two national parks in Texas.',
                       'The park looks totally different from the more populated eastern half of the state.',
                       'For about 118 miles, the Rio Grande River runs through the park',
                       "Big Bend's territory extends to the center of the deepest river channel",
                       'The rest of the land on the other side of the channel belongs to Mexico',
                       "The park's climate reaches extreme temperatures",
                       'This park looks totally different from the more populated eastern half of the state',
                       "It wasn't great."], key=True)
    sheet.render()
