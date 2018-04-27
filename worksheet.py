from docx import Document
from sentence import Sentence
from docx.shared import Pt
from style import POS
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

PUNCT = {',', '.', '-', "'"}


def rindex(lst, val):
    return len(lst) - 1 - lst[::-1].index(val)


class Worksheet:
    def __init__(self, lines=[], title='Sentence Worksheet'):
        self.title = title
        self.font = 'Times New Roman'
        self.lines = [Sentence(line) for line in lines]
        self.doc = Document()

    def render(self, key=True):
        self.doc = Document()
        if key:
            title = self.title + ' (Key)'
        else:
            title = self.title
        self._add_title(title)
        self._add_instructions()
        self._add_title('')
        for line in self.lines:
            self._add_line(line, key=key)
        self.doc.save(title + '.docx')


    def _add_title(self, text):
        title = self.doc.add_paragraph(text)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        font = title.style.font
        font.name = self.font
        font.size = Pt(16)
        font.bold = True

    def _add_instructions(self):
        line = self.doc.add_paragraph()
        font = line.style.font
        font.name = self.font
        font.size = Pt(11)
        label = line.add_run('Label: ')
        self._format_run(label, font_size=11)
        label.bold = True
        self._format_run(line.add_run('Subject, Verb, PN, PA, DO, IO, (prepositional phrase)'), font_size=11)

    def _add_line(self, line, key=True):
        paragraph = self.doc.add_paragraph(style='List Number')
        paragraph.style.font.bold = False
        run = None
        current_prep = -1
        if key:
            pos = line.pos
        else:
            pos = [None] * len(line.pos)
        for index, (word, color) in enumerate(zip(line.doc, pos)):
            if run is not None and str(word) not in PUNCT:
                run = paragraph.add_run(' ')
                self._format_run(run)
            if type(color) == int:
                if index == current_prep:
                    run = paragraph.add_run(str(word) + ')')
                elif index > current_prep:
                    run = paragraph.add_run('(' + str(word))
                    current_prep = rindex(pos, color)
                else:
                    run = paragraph.add_run(str(word))
            else:
                run = paragraph.add_run(str(word))
            self._format_run(run, color=color)

        if str(line.doc[-1]) != '.':
            paragraph.add_run('.')

        self.doc.add_paragraph().style.font.size = Pt(13)


    def _format_run(self, run, color=None, font_size=13):
        if type(color) == POS:
            run.font.color.rgb = color.value
        run.bold = False
        run.font.name = self.font
        run.font.size = Pt(font_size)



if __name__ == '__main__':
    sheet = Worksheet(['In the summer, extreme temperatures of over 100 degrees can give visitors heat strokes.',
                       'Big Bend National Park is one of only two national parks in Texas.',
                       'The park looks totally different from the more populated eastern half of the state.',
                       'For about 118 miles, the Rio Grande River runs through the park',
                       "Big Bend's territory extends to the center of the deepest river channel",
                       'The rest of the land on the other side of the channel belongs to Mexico',
                       "The park's climate reaches extreme temperatures",
                       'This park looks totally different from the more populated eastern half of the state'])
    sheet.render()
