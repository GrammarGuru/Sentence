from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

def create(sentences, filename):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
       
    heading = doc.add_heading("Sentence Worksheet", level = 0)
    for line in sentences:
        paragraph = doc.add_paragraph(line, style = 'ListNumber')
        paragraph.style.paragraph_format.line_spacing = 2
    doc.save(filename)
    
    
#lst = ["I like food.",
#       "I really like food.",
#       "Food is super cool.",
#       "Sentences about food are really fun to read.",
#       "I have nightmares about PPGs.",
#       "PPGs haunt me every day.",
#       "What even is TAG without Mr. Templet.",
#       "Stanford notice my hardcore coding skills."]
#
#create(lst, 'test.docx')